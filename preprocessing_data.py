from docx import Document
import json
import re
import os

""" 
This script was used for data preprocessing.
This script is based on the transcripts as provided to us in a docx format, 
some parts might need to be adjusted when running on transcripts in slightly different formats.
"""

def extract_text(path):
    doc = Document(path)   
    text = []

    # Extract metadata from the interview and save this in csv file
    ID, name, time, respondent, interviewer = get_meta(doc)

    # Get the text from the document
    for para in doc.paragraphs:
        if para.text.strip():  
            text.append(para.text)
    full_text = "\n".join(text)

    # Obtain the transcripts of text from the respondent per timeframe
    full_text = strip(full_text, speaker_label='Z')
    full_text = strip(full_text, speaker_label='Y2')
    full_text = strip(full_text, speaker_label='Y3')
    full_text = strip(full_text, speaker_label=interviewer)

    # Remove time markers like (0:02:35) (0:02:35). and 0:20 - 0:30
    full_text = re.sub(r"\(\d{1,2}:\d{2}:\d{2}\)|\d{1,2}:\d{2} - \d{1,2}:\d{2}", "", full_text)
    full_text = re.sub("\n", " ", full_text)   
    full_text = re.sub("  ", " ", full_text)    
    text_overview = re.sub(rf"({respondent}|Y1):", "ANSWER_", full_text)
    full_text = re.sub(rf"({respondent}|Y1):", "", full_text)

    # Create the directory
    directory = os.path.join("Data", ID)
    os.makedirs(directory, exist_ok=True)  
    
    # Store the overview of transcript
    json_path = os.path.join(directory, "transcript_overview" + '.json')
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(text_overview, f, ensure_ascii=False, indent = 4)

    transcript_by_splits = parse_interview(full_text)

    # Store the transcript per split
    json_path_split = os.path.join(directory, "splits" + '.json')
    with open(json_path_split, "w", encoding="utf-8") as f:
        json.dump(transcript_by_splits, f, ensure_ascii=False, indent=4)


def strip(text, speaker_label):
    # This extracts the speaker's content, their responses and time markers
    pattern = rf'\n?{speaker_label}:\s*.*?(?=\n\w+:|\Z)' 
    def clean_match(match):
        lines = match.group().splitlines()
        cleaned = []
        for line in lines:
            timecodes = re.findall(r'\[\d+\s*min\.?\]', line)
            if timecodes:
                cleaned.extend(timecodes)
        return '\n'.join(cleaned) + '\n' if cleaned else ''
    return re.sub(pattern, clean_match, text, flags=re.DOTALL)

def parse_interview(text):
    text = text.replace("\n", " ")
    time_pattern = re.compile(r'\[(\d+)\s*min\.?\]')
    parts = time_pattern.split(text)
    time_segments = {}
    time_segments[0] = parts[0].strip() 
    for i in range(1, len(parts), 2):
        time = int(parts[i])
        segment = parts[i + 1].strip()
        time_segments[time] = segment
    return time_segments

def get_meta(doc):
    for table in doc.tables:
        for row in table.rows:
            previous_cell = ""
            for cell in row.cells:
                if "is de respondent" in cell.text:
                    a = cell.text
                    respondent = a.replace(" is de respondent", "")
                if "is de interviewer" in cell.text:
                    b = cell.text
                    interviewer = b.replace(" is de interviewer", "")
                if ".mp3" in cell.text.lower() and "_" in cell.text:
                    s = cell.text.replace(".mp3", "").replace(".MP3", "")
                    ID = s.split("_")[0]
                    name = s.split("_")[1]
                    break
                if "Starttijd" in previous_cell:
                    t_s = int(cell.text)
                if "Eindtijd" in previous_cell:
                    t_e = int(cell.text)
                previous_cell = cell.text      
    return ID, name, t_e - t_s, respondent, interviewer
    


# Here, we loop through the different interview transcripts and call the function
# This will create a data directory with a separate directory for each interview, containing its
# splits.json                       per timeframe
# transcript_overview.json          per answer

directory = "Transcripties"                   
for filename in os.listdir(directory):
    print(filename)
    file_path = os.path.join(directory, filename)
    extract_text(file_path)


